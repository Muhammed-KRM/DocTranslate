"""
Core translation system - Tag Formatter V2
Enhanced version with hyperlink preservation, whitespace handling, and per-run formatting
"""
import re
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class TagFormatter:
    """Converts DOCX runs to HTML tags and vice versa - ENHANCED VERSION"""
    
    def paragraph_to_tagged_text(self, para) -> tuple:
        """
        Convert paragraph runs to text with HTML tags
        Returns: (tagged_text, run_styles, hyperlinks)
        """
        tagged_text = ""
        run_styles = []
        hyperlinks = {}
        
        # Iterate over all children (runs and hyperlinks)
        # para.runs ONLY gives top-level runs, skipping hyperlinks!
        child_elements = para._element.xpath('./w:r | ./w:hyperlink')
        
        run_counter = 0
        
        for child in child_elements:
            if child.tag == qn('w:r'):
                # Regular run
                runs = [child]
                link_url = None
            elif child.tag == qn('w:hyperlink'):
                # Hyperlink - can contain multiple runs
                runs = child.xpath('./w:r')
                # Extract URL
                try:
                    r_id = child.get(qn('r:id'))
                    if r_id:
                        link_url = para.part.rels[r_id].target_ref
                    else:
                        link_url = None
                except:
                    link_url = None
            else:
                continue
                
            for run_element in runs:
                # Create a proxy run object for style extraction
                from docx.text.run import Run
                run = Run(run_element, para)
                text = run.text
                
                if not text:
                    continue
                
                # Capture COMPLETE run style
                run_style = {
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': run.font.color.rgb if (run.font.color and run.font.color.rgb) else None,
                    'color_type': run.font.color.type if run.font.color else None, # New: Capture Color Type (AUTO/THEME)
                    'size': run.font.size,
                    'name': run.font.name,
                    'highlight_color': None,
                    'shading': None,
                    'style_id': None # New: Capture Character Style (rStyle)
                }
                
                # Capture properties from the XML element directly to be safe
                try:
                    rPr = run_element.find(qn('w:rPr'))
                    if rPr is not None:
                        # Capture rStyle (Character Style)
                        rStyle = rPr.find(qn('w:rStyle'))
                        if rStyle is not None:
                            run_style['style_id'] = rStyle.get(qn('w:val'))
                            
                        # Capture SHADING (w:shd)
                        shd = rPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            if fill and fill != 'auto':
                                run_style['shading'] = fill
                except:
                    pass
                
                # Capture highlight color (via python-docx)
                try:
                    run_style['highlight_color'] = run.font.highlight_color
                except:
                    pass
                
                run_styles.append(run_style)
                
                # Register hyperlink
                if link_url:
                    hyperlinks[run_counter] = link_url
                
                # Build tags
                tagged_run = f"[R{run_counter}]"
                
                if run.font.bold:
                    text = f"<b>{text}</b>"
                if run.font.italic:
                    text = f"<i>{text}</i>"
                if run.font.underline:
                    text = f"<u>{text}</u>"
                
                if run.font.color and run.font.color.rgb:
                    rgb = str(run.font.color.rgb)
                    text = f'<c rgb="{rgb}">{text}</c>'
                
                try:
                    if run.font.highlight_color:
                        text = f'<h col="{run.font.highlight_color}">{text}</h>'
                except:
                    pass
                
                tagged_run += text + f"[/R{run_counter}]"
                tagged_text += tagged_run
                
                run_counter += 1
        
        return tagged_text, run_styles, hyperlinks
    
    def tags_to_paragraph(self, para, tagged_text: str, run_styles: list, hyperlinks: dict):
        """
        Parse tagged text and rebuild paragraph runs with EXACT formatting
        
        Args:
            para: Paragraph object
            tagged_text: Tagged text with [R0]...[/R0] markers
            run_styles: List of style dicts for each run
            hyperlinks: Dict of {run_idx: url}
        """
        # Clear existing runs
        para.clear()
        
        # Extract runs using markers
        run_pattern = r'\[R(\d+)\](.*?)\[/R\1\]'
        matches = re.findall(run_pattern, tagged_text, re.DOTALL)
        
        for run_idx_str, run_content in matches:
            run_idx = int(run_idx_str)
            
            # Get original style for this run
            if run_idx >= len(run_styles):
                print(f"[WARN] Run index {run_idx} out of range, skipping")
                continue
            
            style = run_styles[run_idx]
            
            # Parse tags within this run
            clean_text = self._parse_and_extract_text(run_content)
            
            if not clean_text:
                continue
            
            # Create run (or hyperlink run)
            if run_idx in hyperlinks:
                run = self._add_hyperlink(para, hyperlinks[run_idx], clean_text)
            else:
                run = para.add_run(clean_text)
            
            # Apply ORIGINAL run style (preserves everything)
            self._apply_run_style(run, style, run_content)
    
    def _parse_and_extract_text(self, tagged_content: str) -> str:
        """Remove all tags from content, keeping only text"""
        # Remove all tags: <b>, </b>, <c rgb="...">, etc.
        # Use regex to strip tags but preserve text
        text = re.sub(r'<c rgb="[^"]*">', '', tagged_content)
        text = re.sub(r'<h col="[^"]*">', '', text)
        text = re.sub(r'<[/]?[biuch]>', '', text)
        
        # Clean HTML entities
        text = text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
        
        return text
    
    def _apply_run_style(self, run, style: dict, tagged_content: str):
        """Apply complete formatting to a run"""
        # Apply base style
        if style.get('name'):
            try:
                run.font.name = style['name']
            except Exception as e:
                print(f"[WARN] Failed to set font name: {e}")
        
        if style.get('size'):
            try:
                run.font.size = style['size']
            except Exception as e:
                print(f"[WARN] Failed to set font size: {e}")
        
        if style.get('color'):
            try:
                run.font.color.rgb = style['color']
            except Exception as e:
                print(f"[WARN] Failed to set font color: {e}")
        elif style.get('color_type') and str(style['color_type']) == 'AUTO (101)':
            # Explicitly set AUTO color to override paragraph style
            try:
                rPr = run._element.get_or_add_rPr()
                color = OxmlElement('w:color')
                color.set(qn('w:val'), 'auto')
                # Remove existing color if any
                existing_color = rPr.find(qn('w:color'))
                if existing_color is not None:
                    rPr.remove(existing_color)
                rPr.append(color)
            except Exception as e:
                print(f"[WARN] Failed to set AUTO color: {e}")
        
        # Apply highlight from ORIGINAL style (not tags)
        if style.get('highlight_color'):
            try:
                run.font.highlight_color = style['highlight_color']
            except Exception as e:
                print(f"[WARN] Failed to set highlight: {e}")
        
        # Parse tags for toggle formatting (bold, italic, underline)
        # Check if tags are present in the tagged content
        if '<b>' in tagged_content:
            run.font.bold = True
        elif style.get('bold'):
            run.font.bold = style['bold']
        
        if '<i>' in tagged_content:
            run.font.italic = True
        elif style.get('italic'):
            run.font.italic = style['italic']
        
        if style.get('underline'):
            run.font.underline = style['underline']
            
        # Apply Character Style (rStyle) via Direct XML Injection
        # We avoid run.style = ... because it does lookup by name, which fails if we only have ID
        if style.get('style_id'):
            try:
                rPr = run._element.get_or_add_rPr()
                # Check if rStyle already exists
                existing_style = rPr.find(qn('w:rStyle'))
                if existing_style is None:
                    # Create new rStyle
                    rStyle = OxmlElement('w:rStyle')
                    rStyle.set(qn('w:val'), style['style_id'])
                    # Insert at beginning of rPr (schema compliance)
                    rPr.insert(0, rStyle)
                else:
                    existing_style.set(qn('w:val'), style['style_id'])
            except Exception as e:
                print(f"[WARN] Failed to force style xml: {e}")
            
        # Apply Shading (Background Color)
        if style.get('shading'):
            try:
                rPr = run._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), style['shading'])
                
                # Remove existing shd if any
                existing_shd = rPr.find(qn('w:shd'))
                if existing_shd is not None:
                    rPr.remove(existing_shd)
                    
                rPr.append(shd)
            except Exception as e:
                print(f"[WARN] Failed to set shading: {e}")
    
    def _get_hyperlink_url(self, run) -> str:
        """Extract hyperlink URL from a run if it exists"""
        try:
            # Check if run is part of a hyperlink
            parent = run._element.getparent()
            if parent.tag == qn('w:hyperlink'):
                # Get relationship ID
                r_id = parent.get(qn('r:id'))
                if r_id:
                    # Get URL from relationship
                    return parent.getparent().part.rels[r_id].target_ref
        except:
            pass
        
        return None
    
    def _add_hyperlink(self, paragraph, url: str, text: str):
        """Add a hyperlink run to paragraph"""
        try:
            # Get the paragraph's part
            part = paragraph.part
            r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
            
            # Create hyperlink element
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('r:id'), r_id)
            
            # Create run within hyperlink
            new_run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            
            # Add run properties (hyperlink style)
            new_run.append(rPr)
            
            # Add text
            text_elem = OxmlElement('w:t')
            text_elem.text = text
            new_run.append(text_elem)
            
            hyperlink.append(new_run)
            paragraph._p.append(hyperlink)
            
            # CRITICAL FIX: Return a proxy Run object for this specific element
            # paragraph.runs DOES NOT include runs inside hyperlinks, so paragraph.runs[-1] is WRONG
            from docx.text.run import Run
            return Run(new_run, paragraph)

        except Exception as e:
            print(f"[WARN] Failed to add hyperlink: {e}")
            # Fallback to regular run
            return paragraph.add_run(text)
    
    def capture_base_style(self, para) -> dict:
        """Capture base style from first run of paragraph"""
        base_style = {}
        
        if para.runs:
            r1 = para.runs[0]
            try:
                base_style['bold'] = r1.font.bold
                base_style['italic'] = r1.font.italic
                base_style['underline'] = r1.font.underline
                base_style['color'] = r1.font.color.rgb if r1.font.color else None
                base_style['size'] = r1.font.size
                base_style['name'] = r1.font.name
                try:
                    base_style['highlight_color'] = r1.font.highlight_color
                except:
                    base_style['highlight_color'] = None
            except:
                pass
        
        return base_style
    
    def should_skip_cell(self, cell) -> bool:
        """Check if table cell should be skipped (symbols, special fonts)"""
        for para in cell.paragraphs:
            for run in para.runs:
                # Skip if Wingdings or similar symbol fonts
                if run.font.name in ["Wingdings", "Symbol", "Webdings"]:
                    return True
                # Skip if only single character (likely a symbol)
                if len(run.text.strip()) == 1 and run.text.strip() in ["✔", "✓", "❌", "✗", "●", "○"]:
                    return True
        
        return False
