"""
Core translation system - Document Translator
Main business logic - environment agnostic
"""
from io import BytesIO
from docx import Document
from .database import DatabaseManager
from .glossary import GlossaryEngine
from .formatters import TagFormatter
from .engines import GoogleTranslator, DeepLTranslator

class DocumentTranslator:
    """
    Main translation engine
    Works with BytesIO streams (environment agnostic)
    """
    
    def __init__(self, db_path: str, engine: str = "google", api_key: str = None):
        """
        Initialize translator
        
        Args:
            db_path: Path to SQLite database
            engine: 'google' or 'deepl'
            api_key: DeepL API key (required if engine='deepl')
        """
        self.db = DatabaseManager(db_path)
        self.glossary = GlossaryEngine(self.db)
        self.formatter = TagFormatter()
        
        # Select translation engine
        if engine == "deepl" and api_key:
            self.engine = DeepLTranslator(api_key)
            print("[INFO] Using DeepL API")
        else:
            self.engine = GoogleTranslator()
            print("[INFO] Using Google Translate")
    
    def translate(self, input_stream: BytesIO, category: str = "general") -> BytesIO:
        """
        Translate a DOCX document
        
        Args:
            input_stream: Input document as BytesIO
            category: Glossary category to use
        
        Returns:
            Translated document as BytesIO
        """
        # Load document
        doc = Document(input_stream)
        
        print(f"[INFO] Loaded document: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        
        # Pass 1: Translate paragraphs
        para_count = self._translate_paragraphs(doc, category)
        print(f"[INFO] Translated {para_count} paragraphs")
        
        # Pass 2: Translate tables
        table_count = self._translate_tables(doc, category)
        print(f"[INFO] Translated {table_count} table cells")
        
        # Pass 3: Final cleanup - catch any remaining Turkish words
        print("[INFO] Running final cleanup pass...")
        cleanup_count = self._final_cleanup_pass(doc, category)
        if cleanup_count > 0:
            print(f"[INFO] Final cleanup translated {cleanup_count} additional items")
        
        # Save to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    
    def _translate_paragraphs(self, doc, category: str) -> int:
        """Translate all paragraphs in document"""
        count = 0
        total = len(doc.paragraphs)
        
        for i, para in enumerate(doc.paragraphs):
            if i % 20 == 0:
                print(f"  Processing paragraph {i}/{total}...", end='\r')
            
            if self._translate_paragraph(para, category):
                count += 1
        
        print()  # New line after progress
        return count
    
    def _translate_paragraph(self, para, category: str) -> bool:
        """
        Translate a single paragraph using tag injection
        Returns True if translation was performed
        """
        import re
        
        # Skip empty paragraphs
        original_text = para.text
        if not original_text or len(original_text.strip()) < 2:
            return False
        
        # Skip if no Turkish characters
        if not self._has_turkish(original_text):
            return False
        
        # New Formatter V2: Get tagged text, styles, and hyperlinks
        # tagged_text will be like "[R0]<b>Text</b>[/R0][R1]Link[/R1]"
        tagged_text, run_styles, hyperlinks = self.formatter.paragraph_to_tagged_text(para)
        
        # SAFETY: If tagging failed, don't proceed
        if not tagged_text or len(tagged_text.strip()) < 2:
            print(f"[WARN] Tagging failed for: {original_text[:50]}")
            return False
        
        # --- PROTECTION PHASE ---
        
        # 1. Whitespace Protection
        # Replace sequences of 2+ spaces with [_WS_N_] to prevent trimming
        protected_text = tagged_text
        if "  " in protected_text:
            protected_text = re.sub(r'( {2,})', lambda m: f' [_WS_{len(m.group(1))}_] ', protected_text)
            
        # 2. Tag Protection
        # Protect [R0], [/R0], <b>, <c>, etc. from translation engine
        # We find anything looking like a tag and hide it
        tag_placeholders = {}
        def protect_tag(match):
            key = f"[_TG_{len(tag_placeholders)}_]"
            tag_placeholders[key] = match.group(0)
            return key
            
        # Matches [R12], [/R12], <c rgb..>, </b> etc.
        tag_pattern = r'\[/?R\d+\]|<[^>]+>'
        protected_text = re.sub(tag_pattern, protect_tag, protected_text)
        
        # 3. Glossary Protection (Existing)
        masked_text, term_placeholders = self.glossary.protect_text(protected_text, category)
        
        # --- TRANSLATION ---
        
        translated_masked = self.engine.translate(masked_text)
        
        # SAFETY: If translation returned nothing, don't delete original!
        if not translated_masked or len(translated_masked.strip()) < 1:
            print(f"[WARN] Translation API returned empty for: {original_text[:50]}")
            return False
        
        # --- RESTORATION PHASE ---
        
        # 1. Restore Glossary Terms
        translated = self.glossary.restore_text(translated_masked, term_placeholders)
        
        # 2. Restore Tags
        # Need to iterate because order might change (though rarely for tags)
        # But we use unique keys, so simple replacement works
        for key, value in tag_placeholders.items():
            translated = translated.replace(key, value)
            
        # 3. Restore Whitespace
        translated = re.sub(r'\s*\[_WS_(\d+)_\]\s*', lambda m: ' ' * int(m.group(1)), translated)
        
        # Fix concatenation errors
        translated = self.glossary.fix_concatenation_errors(translated)
        
        # SAFETY: Final check - ensure we have valid content
        if not translated or len(translated.strip()) < 1:
            print(f"[WARN] Final translated text is empty for: {original_text[:50]}")
            return False
        
        # Check if changed (avoid unnecessary rewrites)
        if translated == tagged_text:
            return False
        
        # Rebuild paragraph with tags and per-run styles
        self.formatter.tags_to_paragraph(para, translated, run_styles, hyperlinks)
        
        return True
    
    def _translate_tables(self, doc, category: str) -> int:
        """Translate all tables in document"""
        count = 0
        
        for table_idx, table in enumerate(doc.tables):
            print(f"  Processing table {table_idx + 1}/{len(doc.tables)}...")
            
            for row in table.rows:
                for cell in row.cells:
                    # Check if cell should be skipped (symbols)
                    if self.formatter.should_skip_cell(cell):
                        continue
                    
                    # Translate cell paragraphs
                    for para in cell.paragraphs:
                        if self._translate_paragraph(para, category):
                            count += 1
        
        return count
    
    def _has_turkish(self, text: str) -> bool:
        """Check if text contains Turkish characters - AGGRESSIVE VERSION"""
        if not text or len(text.strip()) < 1:
            return False
        
        turkish_chars = ['ğ', 'ü', 'ş', 'ı', 'ö', 'ç', 'İ', 'Ğ', 'Ü', 'Ş', 'Ö', 'Ç']
        
        # Turkish specific words (even short ones)
        turkish_words = [
            ' ve ', ' ile ', ' bir ', ' için ', ' olan ', ' olarak ', ' veya ',
            ' gibi ', ' daha ', ' en ', ' çok ', ' az ', ' var ', ' yok ',
            'evet', 'hayır', 'orta', 'tam', 'kriter', 'ay', 'hafta', 'gün',
            'kolay', 'zor', 'hızlı', 'yavaş', 'düşük', 'yüksek', 'minimal',
            'sınırlı', 'kapsamlı', 'genel', 'özel', 'basit', 'karmaşık',
            # User-reported missing words
            'karma', 'otomatik', 'her yerden', 'koordine', 'gerekli', 'bulut',
            'birden fazla', 'temel', 'yok', 'maliyetleri', 'optimize', 'edilir'
        ]
        
        text_lower = text.lower()
        
        # Check for Turkish characters
        if any(char in text for char in turkish_chars):
            return True
        
        # Check for Turkish words
        if any(word in text_lower for word in turkish_words):
            return True
        
        # Check for common Turkish suffixes
        turkish_suffixes = ['ları', 'leri', 'lar', 'ler', 'dır', 'dir', 'tır', 'tir', 'mış', 'miş', 'muş', 'müş']
        if any(text_lower.endswith(suffix) for suffix in turkish_suffixes):
            return True
        
        return False
    
    def _final_cleanup_pass(self, doc, category: str) -> int:
        """
        Final pass to catch any remaining Turkish words that were missed
        This is a safety net to ensure no Turkish remains
        """
        count = 0
        
        # Check all paragraphs again
        for para in doc.paragraphs:
            if para.text and self._has_turkish(para.text):
                # Try to translate again (might have been skipped)
                if self._translate_paragraph(para, category):
                    count += 1
        
        # Check all table cells again
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Don't skip symbols this time if they contain Turkish
                    for para in cell.paragraphs:
                        if para.text and self._has_turkish(para.text):
                            if self._translate_paragraph(para, category):
                                count += 1
        
        return count
    
    def close(self):
        """Cleanup resources"""
        self.db.close()
